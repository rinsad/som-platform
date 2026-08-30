from pathlib import Path
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parent

ACCENT = RGBColor(46, 116, 181)
ACCENT_DARK = RGBColor(31, 77, 120)
TEXT = RGBColor(32, 32, 32)
MUTED = RGBColor(90, 90, 90)
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
BORDER = "B8C4D1"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, color=TEXT, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(int(w * 1440) for w in widths)))

    tbl_grid = tbl.tblGrid
    if tbl_grid is not None:
        tbl.remove(tbl_grid)
    tbl_grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        tbl_grid.append(col)
    tbl.insert(1, tbl_grid)


def strip_inline_code(text):
    return re.sub(r"`([^`]+)`", r"\1", text)


def add_mixed_paragraph(doc, text, style=None, bullet=False, number=False):
    if bullet:
        p = doc.add_paragraph(style="List Bullet")
    elif number:
        p = doc.add_paragraph(style="List Number")
    else:
        p = doc.add_paragraph(style=style)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6 if not bullet and not number else 4)

    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = ACCENT_DARK
        else:
            run = p.add_run(part)
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.font.color.rgb = TEXT
    return p


def add_note_box(doc, label, value):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    set_table_borders(table, "D9E2EC")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_FILL)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + ": ")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.font.color.rgb = ACCENT_DARK
    r = p.add_run(value)
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.font.color.rgb = TEXT
    doc.add_paragraph()


def table_widths(headers):
    n = len(headers)
    lower = [h.lower() for h in headers]
    if n == 2:
        return [2.25, 4.25]
    if n == 3:
        return [2.25, 2.15, 2.10]
    if n == 4:
        if lower == ["user", "email", "role", "what to test"]:
            return [1.28, 1.70, 1.18, 2.34]
        return [1.45, 1.85, 1.25, 1.95]
    return [6.5 / n] * n


def add_table(doc, rows):
    if not rows:
        return
    headers = rows[0]
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_borders(table)
    set_table_width(table, table_widths(headers))
    hdr = table.rows[0]
    for idx, text in enumerate(headers):
        set_cell_text(hdr.cells[idx], strip_inline_code(text), bold=True, color=ACCENT_DARK, size=8.5)
        set_cell_shading(hdr.cells[idx], HEADER_FILL)

    for row_idx, row in enumerate(rows[1:], start=1):
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            set_cell_text(cells[idx], strip_inline_code(text), size=8.5)
            if row_idx % 2 == 0:
                set_cell_shading(cells[idx], "FAFBFC")
    doc.add_paragraph()


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, ACCENT, 18, 10),
        ("Heading 2", 13, ACCENT, 14, 7),
        ("Heading 3", 12, ACCENT_DARK, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_footer(section, title):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(title)
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED


def add_cover(doc, title, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = ACCENT_DARK

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(subtitle)
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = MUTED

    add_note_box(doc, "Audience", "Client end users and testers")
    add_note_box(doc, "Default demo password", "Test@1234")


def parse_table(lines, idx):
    rows = []
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        line = lines[idx].strip()
        parts = [p.strip() for p in line.strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", p or "") for p in parts):
            rows.append(parts)
        idx += 1
    return rows, idx


def build_docx(md_path, out_path, subtitle):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    setup_styles(doc)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    title = lines[0].lstrip("# ").strip()
    add_footer(section, title)
    add_cover(doc, title, subtitle)

    i = 1
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        if not line.strip():
            i += 1
            continue
        if line.strip().startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif re.match(r"^\d+\.\s+", line.strip()):
            add_mixed_paragraph(doc, re.sub(r"^\d+\.\s+", "", line.strip()), number=True)
        elif line.lstrip().startswith("- "):
            add_mixed_paragraph(doc, line.lstrip()[2:], bullet=True)
        else:
            text = line.strip()
            if text in ("Steps:", "Expected result:", "Login as:", "Prerequisite:", "Then login as requester:", "Default password:", "Important note:", "Common approver users:", "Expected final result:", "Conditional workflow:", "Standard workflow:", "Example:"):
                p = add_mixed_paragraph(doc, text)
                p.runs[0].bold = True
                p.runs[0].font.color.rgb = ACCENT_DARK
            else:
                add_mixed_paragraph(doc, text)
        i += 1

    doc.save(out_path)


def main():
    targets = [
        (
            BASE / "capex-end-user-testing-manual.md",
            BASE / "capex-end-user-testing-manual.docx",
            "Self-testing guide for CAPEX request, approval, lifecycle, closure, and audit workflows.",
        ),
        (
            BASE / "pr-end-user-testing-manual.md",
            BASE / "pr-end-user-testing-manual.docx",
            "Self-testing guide for Purchase Request creation, supplier quotations, DoA approvals, and audit workflows.",
        ),
    ]
    for md, out, subtitle in targets:
        build_docx(md, out, subtitle)
        print(out)


if __name__ == "__main__":
    sys.exit(main())
